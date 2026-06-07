# Copyright (c) 2026 Nikolaj Storm Petersen. Licensed under CC BY-NC 4.0.
# Non-commercial use only. If you use or adapt this code, please cite the author.
# See LICENSE and CITATION.cff  |  https://creativecommons.org/licenses/by-nc/4.0/

# ============================================================================
#  create_report.py
#  Stage: Exploratory
#
#  PURPOSE
#    Generates a static Word (.docx) technical report summarising the sepsis
#    prediction model optimisation work (diagnostic findings, applied
#    optimisations, and a hardcoded 1-hour results comparison table). All
#    numbers are literals embedded in this script, not computed from data.
#
#  INPUTS
#    none
#  OUTPUTS
#    /PATH/TO/PROJECT/technical/Results/Sepsis_Model_Optimisation_Report.docx
#
#  USER-EDITABLE SETTINGS  (grep the body for the tag  EDIT:  to find each)
#    save_path                  -  absolute path for the generated .docx report.
#    results table              -  the hardcoded (model, AUROC, recall,
#                                  precision) rows printed into the table.
#
#  REQUIRES: python-docx
# ============================================================================
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('Technical Report: Sepsis Prediction Model Optimisation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "This document details the technical enhancements and corrections applied to the sepsis prediction "
        "modelling pipeline (V13b). The primary goal was to move from a basic training configuration to a "
        "clinically optimized setup, maximizing the predictive power of the engineered feature set."
    )

    # Identified Issues
    doc.add_heading('2. Diagnostic Findings', level=1)
    doc.add_paragraph(
        "Initial analysis of the 'Full Feature' runs revealed that models were plateauing below their theoretical "
        "potential. Three critical bottlenecks were identified:"
    )
    list_items = [
        "Incorrect Class Weighting: The XGBoost 'scale_pos_weight' was hardcoded to 10:1, while the actual prevalence in the 1-hour dataset is ~58:1. This caused the model to significantly under-penalize missed sepsis cases (False Negatives).",
        "Under-training: Models were using conservative estimator counts (100) and high learning rates, leading to early convergence on sub-optimal patterns.",
        "Static Thresholding: Performance was only evaluated at the default 0.5 threshold, which is rarely optimal for imbalanced clinical data where recall is the priority."
    ]
    for item in list_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    # Applied Optimisations
    doc.add_heading('3. Applied Optimisations', level=1)

    doc.add_heading('3.1 Dynamic Imbalance Correction', level=2)
    doc.add_paragraph(
        "We implemented automated class-ratio computation. For every resolution (15-min, 1-hour, 4-hour), the model now "
        "calculates the exact Negative/Positive ratio from the training partition and sets 'scale_pos_weight' accordingly. "
        "This ensures the model treats catching a sepsis case as significantly more important than avoiding a false alarm."
    )

    doc.add_heading('3.2 Feature Selection & Importance Pruning', level=2)
    doc.add_paragraph(
        "To reduce noise and prevent overfitting, we implemented an automated feature selector. The pipeline trains an "
        "initial XGBoost model to rank all 95 engineered features. It then evaluates top-N subsets (20, 35, 50, and Full). "
        "For the V13b dataset, the full feature set with optimized training was found to be the most robust."
    )

    doc.add_heading('3.3 Clinical Threshold Tuning (F2-Score Focus)', level=2)
    doc.add_paragraph(
        "Because sepsis detection is a 'must-not-miss' task, we implemented F2-Score optimization. Unlike the standard F1-score, "
        "the F2-score weights recall twice as heavily as precision. The script sweeps decision thresholds (from 0.05 to 0.50) "
        "to find the point that provides the best clinical balance."
    )

    # Results
    doc.add_heading('4. Results Comparison (1-Hour Resolution)', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model Version'
    hdr_cells[1].text = 'AUROC'
    hdr_cells[2].text = 'Recall'
    hdr_cells[3].text = 'Precision'

    # EDIT: results table - hardcoded (model, AUROC, recall, precision) rows
    results = [
        ('Baseline XGBoost', '0.763', '6.9%', '0.123'),
        ('Optimised XGBoost (Default)', '0.851', '77.7%', '0.053'),
        ('Optimised XGBoost (F2-Opt)', '0.851', '95.1%', '0.034'),
        ('Baseline Random Forest', '0.745', '61.3%', '0.039'),
        ('Optimised Random Forest', '0.828', '66.2%', '0.052')
    ]

    for mv, auroc, recall, prec in results:
        row_cells = table.add_row().cells
        row_cells[0].text = mv
        row_cells[1].text = auroc
        row_cells[2].text = recall
        row_cells[3].text = prec

    doc.add_paragraph(
        "\nConclusion: The optimization phase yielded a +0.088 gain in AUROC for XGBoost and enabled a massive increase in sensitivity "
        "(from ~7% to >95%) by correctly managing the precision-recall trade-off."
    )

    # Save
    # EDIT: save_path - absolute path for the generated .docx report
    save_path = '/PATH/TO/PROJECT/technical/Results/Sepsis_Model_Optimisation_Report.docx'
    doc.save(save_path)
    return save_path

if __name__ == "__main__":
    path = create_report()
    print(f"✅ Report saved to: {path}")
