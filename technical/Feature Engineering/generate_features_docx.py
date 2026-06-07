# Copyright (c) 2026 Nikolaj Storm Petersen. Licensed under CC BY-NC 4.0.
# Non-commercial use only. If you use or adapt this code, please cite the author.
# See LICENSE and CITATION.cff  |  https://creativecommons.org/licenses/by-nc/4.0/

# ============================================================================
#  generate_features_docx.py
#  Stage: 3 - Feature Engineering (documentation generator)
#
#  PURPOSE
#    Generates a Word (.docx) reference document describing every variable
#    and engineered feature used in the V13b pipeline, split into
#    ground-truth target variables, SOFA inputs, demographics, raw vitals,
#    derived clinical scores, and advanced temporal kinematics.
#
#  INPUTS
#    none (all content is hard-coded in this script)
#  OUTPUTS
#    /PATH/TO/PROJECT/technical/MIMIC_V13b_Data_Features.docx
#
#  USER-EDITABLE SETTINGS  (grep the body for the tag  EDIT:  to find each)
#    output docx path  -  where the generated .docx is written
#
#  REQUIRES: python-docx
# ============================================================================
import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

doc = Document()

# Add a Title
title = doc.add_heading('MIMIC-IV Data Features for V13b Pipeline', 0)

# Add intro
doc.add_paragraph('This document outlines all of the variables and data features extracted from the MIMIC-IV database for the V13b pipeline, including the ground-truth target variables (like those for the SOFA score) not passed as training features.')

# Section 1
doc.add_heading('1. Variables Extracted Exclusively to Define Ground-Truth Targets (Not Fed to Model)', level=2)
doc.add_paragraph('These variables are extracted to determine physiological severity and exact clinical onset ("Time Zero"), but are expressly stripped from the training dataset to prevent data leakage.')

table1 = doc.add_table(rows=1, cols=4)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'MIMIC-IV Source Logic'
hdr_cells[3].text = 'Role in V13b'

data1 = [
    ('Suspected Infection Time', 'The proximal overlap of administering antibiotics alongside drawing a blood culture (within a 72h window).', 'Pre-computed mimic_derived.sepsis3 table.', 'Anchors the prediction window (e.g., used to define the 6-hour and 12-hour predictive horizons).'),
    ('SOFA Score Delta (≥ 2)', 'Sequential Organ Failure Assessment. Consists of lab tests & clinical signs: PaO2/FiO2, GCS, MAP/Vasopressors, Bilirubin, Platelets, Creatinine/Urine tracking.', 'Pre-computed mimic_derived.sepsis3 table.', 'Defines the absolute severity threshold for the Sepsis-3 label.'),
    ('ICD Sepsis Etiologies', 'Categorical diagnostic billing codes documenting localized etiology origins.', 'Mapped ICD-9 and ICD-10 codes.', 'Used to define the independent multi-target sub-streams (target_resp, target_uri, target_other).'),
    ('Sepsis Onset Timestamp', 'The exact timestamp where the SOFA ≥ 2 criteria aligns with infection suspicion.', 'Derived temporal intersection.', 'Acts as the operational cutoff limit; vitals after this point are dropped from training.')
]
for item in data1:
    row_cells = table1.add_row().cells
    row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = item

# Section 2
doc.add_heading('2. Raw Variables Extracted to Calculate the SOFA Score (Target Generation)', level=2)
doc.add_paragraph('The SOFA score algorithm in MIMIC-IV requires querying several underlying tables to evaluate the 6 target organ systems. These are essential for building the absolute, unbiased ground-truth labels. While the predictive models are designed to be entirely "lab-free" and rely purely on high-frequency vitals, the lab events and medication logs are fundamental in defining exactly when and if the patient became severely injured by an infection.')

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Organ System'
hdr_cells2[1].text = 'Specific MIMIC-IV Variables Pulled'
hdr_cells2[2].text = 'Data Source/Events in MIMIC'
hdr_cells2[3].text = 'Purpose in SOFA Calculation'

data2 = [
    ('Respiration', 'PaO2 (Arterial Oxygen Partial Pressure)\nFiO2 (Fraction of Inspired Oxygen)\nVentilation Status', 'labevents (Blood Gas)\nchartevents (Vents)', 'Used to calculate the P/F Ratio (PaO2/FiO2). Lower ratios indicate higher respiratory failure.'),
    ('Coagulation', 'Platelets (Platelet Count)', 'labevents', 'Dropping platelet counts act as the primary marker for coagulation failure.'),
    ('Liver', 'Bilirubin, Total', 'labevents', 'Elevated bilirubin levels precisely measure hepatic (liver) failure/dysfunction.'),
    ('Cardiovascular', 'Mean Arterial Pressure (MAP)\nDopamine (Rate/Dose)\nEpinephrine (Rate/Dose)\nNorepinephrine (Rate/Dose)\nDobutamine (Rate/Dose)', 'chartevents (MAP)\ninputevents (Vasopressors)', 'MAP < 70 mmHg flags initial cardiovascular failure. Vasopressors indicate severe shock.'),
    ('Central Nervous System', 'Glasgow Coma Scale (GCS) - Motor\nGCS - Verbal\nGCS - Eyes', 'chartevents', 'A dropping GCS measures neurological degradation and brain dysfunction.'),
    ('Renal (Kidneys)', 'Creatinine (Blood level)\nUrine Output (Cumulative volume)', 'labevents (Creatinine)\noutputevents (Urine)', 'High creatinine or severely depressed urine output indicates acute kidney injury (AKI).')
]

for item in data2:
    row_cells = table2.add_row().cells
    row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = item

doc.add_heading('Additional Operational Variables (For the Suspicion of Infection Anchors)', level=3)
table2a = doc.add_table(rows=1, cols=4)
table2a.style = 'Table Grid'
hdr_cells2a = table2a.rows[0].cells
hdr_cells2a[0].text = 'Variable'
hdr_cells2a[1].text = 'Specific MIMIC-IV Variables Pulled'
hdr_cells2a[2].text = 'Data Source in MIMIC'
hdr_cells2a[3].text = 'Purpose'

data2a = [
    ('Antibiotics', 'Itemid / Route / Start & End times', 'pharmacy, prescriptions, inputevents', 'To identify when therapeutic intervention for an infection began.'),
    ('Blood Cultures', 'Specimen ID, Charttime, Specimen Type', 'microbiologyevents', 'To identify when physicians officially ordered a culture to test for systemic infection.')
]

for item in data2a:
    row_cells = table2a.add_row().cells
    row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = item

# Section 3
doc.add_heading('3. Static Demographics & Context (Exposed to Model)', level=2)
table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Variable'
hdr_cells3[1].text = 'Description'
hdr_cells3[2].text = 'MIMIC-IV Source Logic'
hdr_cells3[3].text = 'Role in V13b'

data3 = [
    ('Age', "Patient's age at the time of admission.", 'patients.anchor_age + Admission Delta.', 'Anchors baseline physiological expectations.'),
    ('Weight (kg)', "Patient's admission or daily recorded weight.", 'chartevents (IDs: 226512, 224639).', 'Used to scale systemic physiological responses.'),
    ('Time Since Admission', 'Cumulative temporal count tracking hours since the absolute ICU intime.', 'TIMESTAMP_DIFF(charttime, intime).', 'Provides sequence awareness to the independent model trees.')
]

for item in data3:
    row_cells = table3.add_row().cells
    row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = item

# Section 4
doc.add_heading('4. Primary Real-Time Vital Sensors (Exposed to Model)', level=2)
table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Table Grid'
hdr_cells4 = table4.rows[0].cells
hdr_cells4[0].text = 'Variable'
hdr_cells4[1].text = 'Description'
hdr_cells4[2].text = 'MIMIC-IV Source Logic (chartevents)'

data4 = [
    ('Heart Rate', 'Cardiac rate sequence.', 'IDs: 220045'),
    ('Respiratory Rate', 'Mechanical or spontaneous breath rate.', 'IDs: 220210, 224690'),
    ('SpO2', 'Peripheral oxygen saturation percentage.', 'IDs: 220277'),
    ('Temperature', 'Thermometric tracking (converted strictly to Celsius).', 'IDs: 223761, 223762'),
    ('Systolic BP (SBP)', 'Maximum arterial exertion.', 'IDs: 220179, 220050'),
    ('Diastolic BP (DBP)', 'Minimum arterial exertion.', 'IDs: 220180, 220051')
]

for item in data4:
    row_cells = table4.add_row().cells
    row_cells[0].text, row_cells[1].text, row_cells[2].text = item

# Section 5
doc.add_heading('5. Derived Clinical Scores (Exposed to Model)', level=2)
table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Table Grid'
hdr_cells5 = table5.rows[0].cells
hdr_cells5[0].text = 'Variable'
hdr_cells5[1].text = 'Description'
hdr_cells5[2].text = 'Formula / Logic'

data5 = [
    ('MAP', 'Mean Arterial Pressure.', '(SBP + 2 * DBP) / 3'),
    ('Shock Index', 'Indicator of occult hemorrhagic or septic shock.', 'Heart Rate / SBP'),
    ('NEWS Score', 'National Early Warning Score (sensor features only).', 'Point-based aggregation of deviation in vitals.'),
    ('Partial qSOFA', 'Simplified early-warning logic.', 'Flags given mechanically if SBP <= 100 or RespRate >= 22.'),
    ('Pulse Pressure', 'Rigidity or stroke volume indicator.', 'SBP - DBP'),
    ('Rate Pressure Product (RPP)', 'Myocardial oxygen consumption index.', 'Heart Rate * SBP'),
    ('Modified Shock Index (MSI)', 'Shock variance scaled against MAP.', 'Heart Rate / MAP'),
    ('Fever-Tachycardia', 'Interactive flag targeting Uri/Infection spikes.', 'Heart Rate * max(Temp - 37, 0)'),
    ('Respiratory Distress', 'Nonlinear penalty for fast breathing with low sat.', 'RespRate * (100 - SpO2)')
]

for item in data5:
    row_cells = table5.add_row().cells
    row_cells[0].text, row_cells[1].text, row_cells[2].text = item

# Section 6
doc.add_heading('6. Advanced Mathematical & Temporal Kinematics (Exposed to Model)', level=2)
table6 = doc.add_table(rows=1, cols=3)
table6.style = 'Table Grid'
hdr_cells6 = table6.rows[0].cells
hdr_cells6[0].text = 'Feature Pattern'
hdr_cells6[1].text = 'Variables Created'
hdr_cells6[2].text = 'Purpose in Architecture'

data6 = [
    ('Acceleration (accel_X)', 'accel_heart_rate, accel_resprate, accel_spo2 (etc. for all vitals)', 'Captures the absolute second derivative (the rate of change of the rate of change). Detects non-linear crashes.'),
    ('CUSUM Offsets', 'cusum_pos_X, cusum_neg_X (applied to Temp, SBP, DBP)', 'Cumulative Sum logic tracks sustained structural deviation away from the patient personal historical baseline.'),
    ('4H Slope Tracers', 'slope_4h_temp, slope_4h_sbp, slope_4h_dbp', 'Statically measures the overarching trajectory vector over a rolling 4-hour window, bypassing immediate noise.'),
    ('Baseline Statistics', 'exp_std_X, X_sd, exp_max_X, exp_min_X', 'Rolling volatility monitors. Extracts the statistical unreliability/arrhythmic nature of a sensor feed.'),
    ('Score Deltas', 'news_delta, qsofa_delta', 'Directly tracks if traditional medical scores are worsening sequentially.')
]

for item in data6:
    row_cells = table6.add_row().cells
    row_cells[0].text, row_cells[1].text, row_cells[2].text = item

# EDIT: output .docx path
doc.save('/PATH/TO/PROJECT/technical/MIMIC_V13b_Data_Features.docx')
print("Successfully created docx")
